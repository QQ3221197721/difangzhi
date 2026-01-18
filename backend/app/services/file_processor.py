"""
地方志数据智能管理系统 - 文件处理服务
"""
import os
import io
from typing import Dict, Any, Optional

import pandas as pd
from docx import Document as DocxDocument
import PyPDF2
import structlog

logger = structlog.get_logger()


class FileProcessorService:
    """文件处理服务"""
    
    SUPPORTED_TYPES = {
        '.pdf': 'pdf',
        '.doc': 'doc',
        '.docx': 'docx',
        '.xls': 'excel',
        '.xlsx': 'excel',
        '.txt': 'text',
        '.csv': 'csv'
    }
    
    async def extract_content(self, file_path: str, content: bytes) -> Dict[str, Any]:
        """从文件中提取内容"""
        ext = os.path.splitext(file_path)[1].lower()
        file_type = self.SUPPORTED_TYPES.get(ext)
        
        if not file_type:
            return {"error": f"Unsupported file type: {ext}"}
        
        try:
            if file_type == 'pdf':
                return await self._extract_pdf(content)
            elif file_type == 'docx':
                return await self._extract_docx(content)
            elif file_type == 'excel':
                return await self._extract_excel(content)
            elif file_type == 'text':
                return await self._extract_text(content)
            elif file_type == 'csv':
                return await self._extract_csv(content)
            else:
                return {"error": f"Processor not implemented for: {file_type}"}
        except Exception as e:
            logger.error(f"File extraction error: {e}")
            return {"error": str(e)}
    
    async def _extract_pdf(self, content: bytes) -> Dict[str, Any]:
        """提取 PDF 内容"""
        try:
            pdf_file = io.BytesIO(content)
            reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            full_text = "\n".join(text_parts)
            
            return {
                "content": full_text[:2000] if len(full_text) > 2000 else full_text,
                "full_text": full_text,
                "pages": len(reader.pages),
                "metadata": dict(reader.metadata) if reader.metadata else {}
            }
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return {"error": str(e)}
    
    async def _extract_docx(self, content: bytes) -> Dict[str, Any]:
        """提取 Word 文档内容"""
        try:
            doc_file = io.BytesIO(content)
            doc = DocxDocument(doc_file)
            
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            full_text = "\n".join(paragraphs)
            
            # 提取表格
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                "content": full_text[:2000] if len(full_text) > 2000 else full_text,
                "full_text": full_text,
                "paragraphs": len(paragraphs),
                "tables": tables[:5]  # 最多返回5个表格
            }
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return {"error": str(e)}
    
    async def _extract_excel(self, content: bytes) -> Dict[str, Any]:
        """提取 Excel 内容"""
        try:
            excel_file = io.BytesIO(content)
            dfs = pd.read_excel(excel_file, sheet_name=None)
            
            sheets_data = {}
            full_text_parts = []
            
            for sheet_name, df in dfs.items():
                # 转换为文本
                sheet_text = df.to_string()
                full_text_parts.append(f"[{sheet_name}]\n{sheet_text}")
                
                # 保存摘要数据
                sheets_data[sheet_name] = {
                    "rows": len(df),
                    "columns": list(df.columns),
                    "preview": df.head(10).to_dict('records')
                }
            
            full_text = "\n\n".join(full_text_parts)
            
            return {
                "content": full_text[:2000] if len(full_text) > 2000 else full_text,
                "full_text": full_text,
                "sheets": sheets_data
            }
        except Exception as e:
            logger.error(f"Excel extraction error: {e}")
            return {"error": str(e)}
    
    async def _extract_text(self, content: bytes) -> Dict[str, Any]:
        """提取纯文本内容"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
            text = None
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                text = content.decode('utf-8', errors='ignore')
            
            return {
                "content": text[:2000] if len(text) > 2000 else text,
                "full_text": text,
                "length": len(text)
            }
        except Exception as e:
            logger.error(f"Text extraction error: {e}")
            return {"error": str(e)}
    
    async def _extract_csv(self, content: bytes) -> Dict[str, Any]:
        """提取 CSV 内容"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312']
            df = None
            
            for encoding in encodings:
                try:
                    csv_file = io.BytesIO(content)
                    df = pd.read_csv(csv_file, encoding=encoding)
                    break
                except (UnicodeDecodeError, pd.errors.ParserError):
                    continue
            
            if df is None:
                return {"error": "Failed to parse CSV file"}
            
            full_text = df.to_string()
            
            return {
                "content": full_text[:2000] if len(full_text) > 2000 else full_text,
                "full_text": full_text,
                "rows": len(df),
                "columns": list(df.columns),
                "preview": df.head(10).to_dict('records')
            }
        except Exception as e:
            logger.error(f"CSV extraction error: {e}")
            return {"error": str(e)}


# 单例实例
file_processor = FileProcessorService()
