"""
文件处理服务 - 支持PDF、TXT、DOC文件解析
"""
import os
from typing import Optional
import aiofiles
from loguru import logger


class FileProcessorService:
    """文件处理服务"""
    
    async def extract_text(self, file_path: str, file_type: str) -> str:
        """从文件中提取文本内容"""
        try:
            if file_type == 'txt':
                return await self._extract_from_txt(file_path)
            elif file_type == 'pdf':
                return await self._extract_from_pdf(file_path)
            elif file_type in ['doc', 'docx']:
                return await self._extract_from_doc(file_path)
            else:
                raise ValueError(f"不支持的文件类型: {file_type}")
        except Exception as e:
            logger.error(f"提取文本失败: {e}")
            raise
    
    async def _extract_from_txt(self, file_path: str) -> str:
        """从TXT文件提取文本"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
        
        for encoding in encodings:
            try:
                async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
                    content = await f.read()
                    return content
            except UnicodeDecodeError:
                continue
        
        raise ValueError("无法识别文件编码")
    
    async def _extract_from_pdf(self, file_path: str) -> str:
        """从PDF文件提取文本"""
        import PyPDF2
        
        text_parts = []
        
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        return '\n'.join(text_parts)
    
    async def _extract_from_doc(self, file_path: str) -> str:
        """从Word文档提取文本"""
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return '\n'.join(text_parts)
    
    def get_file_info(self, file_path: str) -> dict:
        """获取文件信息"""
        stat = os.stat(file_path)
        return {
            "size": stat.st_size,
            "created": stat.st_ctime,
            "modified": stat.st_mtime
        }
